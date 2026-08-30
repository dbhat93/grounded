"""Generate mock sales knowledge docs for the ingestion layer.

All content is fictional (the Kestrel testbed). Two things are seeded on
purpose so the pipeline has something to catch:
  - a CROSS-SOURCE CONFLICT: the capability matrix (a Google Sheet) says EU
    data residency is Roadmap; the security posture (a SharePoint doc) says it
    is GA. Same topic, two sources, disagreeing.
  - a STALE doc: the SentinelIQ battle card was last verified months ago.

Run: python3 fixtures/build_fixtures.py
"""
import os

HERE = os.path.dirname(os.path.abspath(__file__))


def build_capability_matrix():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Capabilities"
    ws.append(["Capability", "Status", "Detail", "Last Verified"])
    rows = [
        ["Real-time transaction scoring", "GA", "Scored in real time, p95 under 200ms.", "2026-06-14"],
        ["AML transaction monitoring", "GA", "Configurable typologies and thresholds.", "2026-06-14"],
        ["Check fraud detection", "Beta", "Deposit and image analysis, limited customers.", "2026-07-02"],
        ["EU data residency", "Roadmap", "On the roadmap, no committed date; US-only today.", "2026-07-18"],
        ["Consortium shared signals", "Roadmap", "Targeted H1 2027; does not exist today.", "2026-07-18"],
        ["Automated SAR filing", "Not supported", "Kestrel prepares narratives; a human files.", "2026-06-14"],
        # net-new facts (not in the hand-authored KB) so promotion has real gaps to fill
        ["ServiceNow case sync", "GA", "Native ServiceNow integration for syncing fraud cases and alert status.", "2026-06-20"],
        ["Databricks export", "GA", "Scheduled export of alerts and decisions to a Databricks lakehouse.", "2026-06-20"],
    ]
    for r in rows:
        ws.append(r)
    wb.save(os.path.join(HERE, "capability_matrix.xlsx"))


def build_security_questionnaire():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "SIG"
    ws.append(["Question", "Answer", "Status", "Last Verified"])
    rows = [
        ["Are you SOC 2 Type II certified?",
         "Yes. SOC 2 Type II report, audited annually, available under NDA.", "GA", "2026-04-11"],
        ["Are you PCI DSS compliant?",
         "Assessed as a PCI DSS Level 1 service provider; AOC under NDA.", "GA", "2026-04-11"],
        ["How is data encrypted?",
         "AES-256 at rest, TLS 1.2+ in transit, AWS KMS for keys.", "GA", "2026-04-11"],
        ["Do you support SSO and SCIM?",
         "SAML 2.0 and OIDC SSO are GA; SCIM provisioning is in beta.", "GA", "2026-05-30"],
    ]
    for r in rows:
        ws.append(r)
    wb.save(os.path.join(HERE, "security_questionnaire.xlsx"))


def build_security_posture():
    from docx import Document
    d = Document()
    d.add_heading("Kestrel Security Posture", 0)
    sections = [
        ("SOC 2", "Kestrel maintains a SOC 2 Type II report, audited annually by a third party."),
        ("Data hosting", "Kestrel runs on AWS in US regions (us-east-1 and us-west-2)."),
        # CONFLICT: matrix says EU data residency is Roadmap; this says GA.
        ("EU data residency", "EU data residency is generally available for enterprise customers."),
        ("Encryption", "Data is encrypted at rest with AES-256 and in transit with TLS 1.2 or higher."),
    ]
    for heading, body in sections:
        d.add_heading(heading, level=1)
        d.add_paragraph(body)
    d.add_paragraph("Document last reviewed: 2026-07-25")
    d.save(os.path.join(HERE, "security_posture.docx"))


def build_pricing_guide():
    from docx import Document
    d = Document()
    d.add_heading("Kestrel Pricing Guide (Internal)", 0)
    d.add_heading("Pricing model", level=1)
    d.add_paragraph("Annual platform fee plus a per-decision usage component. "
                    "Specific numbers are deal-specific; engage the account team for a quote.")
    d.add_heading("Free trial", level=1)
    d.add_paragraph("No self-serve free trial. A scoped proof-of-concept is offered as part "
                    "of the evaluation.")
    d.add_paragraph("Document last reviewed: 2026-05-30")
    d.save(os.path.join(HERE, "pricing_guide.docx"))


def build_battle_card():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 10, "Competitive Battle Card: SentinelIQ",
                   new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font("Helvetica", size=11)
    body = (
        "Status: Battle card\n"
        "Last verified: 2026-01-15\n\n"
        "SentinelIQ is strong on rules-based AML but has no sub-200ms real-time ML "
        "scoring; Kestrel real-time scoring is GA. Both are SOC 2 Type II. SentinelIQ "
        "has no Jack Henry Symitar integration; Kestrel does.\n\n"
        "Position on real-time ML and credit-union core coverage. Verify any SentinelIQ "
        "capability claim against their current materials before stating it to the buyer."
    )
    for line in body.split("\n"):
        pdf.set_x(pdf.l_margin)
        if line.strip():
            pdf.multi_cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.ln(4)
    pdf.output(os.path.join(HERE, "battle_card_sentineliq.pdf"))


if __name__ == "__main__":
    build_capability_matrix()
    build_security_questionnaire()
    build_security_posture()
    build_pricing_guide()
    build_battle_card()
    print("fixtures written to", HERE)
    for f in sorted(os.listdir(HERE)):
        if not f.endswith(".py"):
            print("  ", f)
