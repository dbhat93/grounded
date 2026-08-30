"""Format extraction: RawDoc bytes -> structural ExtractedUnits with locators.

Real extractors for xlsx (openpyxl), docx (python-docx), pdf (pypdf), and plain
transcripts. Each unit carries a precise locator so the fact it becomes can cite
exactly where it came from.
"""
import re

from .model import ExtractedUnit

_DATE_RE = re.compile(r"(\d{4}-\d{2}-\d{2})")
_LINE_RE = re.compile(r"^\s*(?:\[([^\]]+)\]\s*)?([A-Za-z][\w ]*?):\s*(.*)$")


def extract(raw):
    fn = _EXTRACTORS.get(raw.fmt)
    if fn is None:
        return []
    return fn(raw)


def _xlsx(raw):
    from openpyxl import load_workbook
    wb = load_workbook(raw.path, read_only=True, data_only=True)
    units = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h).strip() if h is not None else "" for h in rows[0]]
        for n, row in enumerate(rows[1:], start=2):
            cells = ["" if c is None else str(c) for c in row]
            if not any(c.strip() for c in cells):
                continue
            fields = {h: v for h, v in zip(headers, cells) if h}
            units.append(ExtractedUnit(
                locator="%s!row%d" % (ws.title, n),
                text=" | ".join(c for c in cells if c),
                fields=fields))
    return units


def _docx(raw):
    from docx import Document
    doc = Document(raw.path)
    reviewed = ""
    for p in doc.paragraphs:
        m = re.search(r"last reviewed:\s*" + _DATE_RE.pattern, p.text, re.I)
        if m:
            reviewed = m.group(1)
    units = []
    heading, body = None, []
    def flush():
        if heading is not None and body:
            units.append(ExtractedUnit(
                locator="section: %s" % heading,
                text="\n".join(body),
                fields={"heading": heading, "body": "\n".join(body),
                        "doc_last_reviewed": reviewed}))
    for p in doc.paragraphs:
        style = (p.style.name or "") if p.style else ""
        if style.startswith("Heading") and p.text.strip():
            flush()
            heading, body = p.text.strip(), []
        elif style == "Title":
            continue
        elif p.text.strip() and not p.text.lower().startswith("document last reviewed"):
            body.append(p.text.strip())
    flush()
    return units


def _pdf(raw):
    from pypdf import PdfReader
    reader = PdfReader(raw.path)
    text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    return [ExtractedUnit(
        locator="pages 1-%d" % len(reader.pages),
        text=text,
        fields={"text": text})]


def _transcript(raw):
    units = []
    with open(raw.path, encoding="utf-8") as fh:
        for i, line in enumerate(fh, start=1):
            line = line.rstrip("\n")
            if not line.strip():
                continue
            m = _LINE_RE.match(line)
            if not m:
                continue
            ts, speaker, said = m.group(1), m.group(2).strip(), m.group(3).strip()
            if not said:
                continue
            units.append(ExtractedUnit(
                locator="%s %s" % (ts or ("line%d" % i), speaker),
                text=said,
                fields={"speaker": speaker, "said": said, "ts": ts or ""}))
    return units


_EXTRACTORS = {
    "xlsx": _xlsx,
    "docx": _docx,
    "pdf": _pdf,
    "transcript": _transcript,
}
